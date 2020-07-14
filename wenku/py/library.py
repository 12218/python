#!/usr/bin/env python
# coding: utf-8

# In[32]:


from docx import Document
import requests
import re
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from docx.shared import Pt


# In[24]:


class wenku:
    def __init__(self, url):
        super().__init__()
        self.url = url
        self.doc = ''
        self.header = {'User-Agent': 'Googlebot'}
        
    def get_HTML(self):
        self.response = requests.get(self.url, headers = self.header)
        print(self.response)
        # print(self.response.text)
        
    def get_Doc(self):
        soup = BeautifulSoup(self.response.text, features = 'lxml')
        text = soup.find_all('div', {'class':'doc-reader'})
        print(text[0].get_text())
        return text[0].get_text()


# In[36]:


class write_doc:
    def __init__(self):
        super().__init__()
    def write(self, doc):
        document = Document()
        document.styles["Normal"].font.name = u'宋体'
        document.styles["Normal"].font.size = Pt(14)
        document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        paragraph = document.add_paragraph()
        paragraph.add_run(doc)
        document.save('./artical/text.docx')


# In[38]:


if __name__ == '__main__':
    html = wenku('https://wenku.baidu.com/view/b2072721ccbff121dd3683a2.html')
    artical = write_doc()
    try:
        html.get_HTML()
        if html.response.status_code == 200:
            # html.get_Doc()
            artical.write(html.get_Doc())
            print('Finished!')
        else:
            print('Status Code is not 200!')
    except:
        print('Request Error!')


# In[ ]:




